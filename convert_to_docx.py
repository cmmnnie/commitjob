#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown을 Word 문서(.docx)로 변환하는 스크립트
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def add_heading(doc, text, level):
    """제목 추가"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, style=None):
    """일반 문단 추가"""
    if style:
        p = doc.add_paragraph(text, style=style)
    else:
        p = doc.add_paragraph(text)
    return p

def add_code_block(doc, code_text):
    """코드 블록 추가 (회색 배경)"""
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    # 배경색은 Word에서 직접 지원하지 않으므로 폰트 색상만 변경
    run.font.color.rgb = RGBColor(0, 0, 128)
    return p

def parse_markdown_to_docx(md_file, output_file):
    """Markdown 파일을 읽어 Word 문서로 변환"""

    doc = Document()

    # 기본 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(11)

    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    in_code_block = False
    code_block_content = []
    in_list = False

    i = 0
    while i < len(lines):
        line = lines[i]

        # 코드 블록 시작/끝
        if line.strip().startswith('```'):
            if in_code_block:
                # 코드 블록 끝
                code_text = '\n'.join(code_block_content)
                add_code_block(doc, code_text)
                code_block_content = []
                in_code_block = False
            else:
                # 코드 블록 시작
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # 빈 줄
        if not line.strip():
            if not in_list:
                doc.add_paragraph()
            i += 1
            continue

        # H1 제목 (# )
        if line.startswith('# ') and not line.startswith('## '):
            text = line.replace('# ', '').strip()
            add_heading(doc, text, 1)
            in_list = False

        # H2 제목 (## )
        elif line.startswith('## ') and not line.startswith('### '):
            text = line.replace('## ', '').strip()
            add_heading(doc, text, 2)
            in_list = False

        # H3 제목 (### )
        elif line.startswith('### '):
            text = line.replace('### ', '').strip()
            add_heading(doc, text, 3)
            in_list = False

        # 리스트 항목 (-, ├──, │, └──)
        elif line.strip().startswith(('- ', '├── ', '│   ├── ', '│   └── ', '└── ', '   ├── ', '   └── ')):
            # 리스트 마커 제거
            cleaned = re.sub(r'^[\s│├└─]*', '', line)
            if cleaned.strip():
                p = doc.add_paragraph(cleaned.strip(), style='List Bullet')
                in_list = True

        # 일반 문단
        else:
            # 볼드 처리 (**text**)
            text = line.strip()
            if text:
                p = add_paragraph(doc, text)
                in_list = False

        i += 1

    # 문서 저장
    doc.save(output_file)
    print(f'✅ Word 문서가 생성되었습니다: {output_file}')

if __name__ == '__main__':
    input_file = '기술아키텍처.md'
    output_file = '기술아키텍처.docx'

    try:
        parse_markdown_to_docx(input_file, output_file)
    except Exception as e:
        print(f'❌ 오류 발생: {e}')
        import traceback
        traceback.print_exc()
