#!/usr/bin/env python
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
import re
import sys

def parse_markdown_to_docx(md_file, output_file):
    doc = Document()

    # 기본 스타일
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(11)

    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    in_code_block = False
    code_block_content = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # 코드 블록
        if line.strip().startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_block_content)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 0, 128)
                code_block_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # 빈 줄
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # 제목
        if line.startswith('# ') and not line.startswith('## '):
            text = line.replace('# ', '').strip()
            doc.add_heading(text, 1)
        elif line.startswith('## ') and not line.startswith('### '):
            text = line.replace('## ', '').strip()
            doc.add_heading(text, 2)
        elif line.startswith('### '):
            text = line.replace('### ', '').strip()
            doc.add_heading(text, 3)
        # 리스트
        elif line.strip().startswith(('- ', '├── ', '│', '└── ')):
            cleaned = re.sub(r'^[\s│├└─]*', '', line)
            if cleaned.strip():
                doc.add_paragraph(cleaned.strip(), style='List Bullet')
        # 일반 문단
        else:
            text = line.strip()
            if text:
                doc.add_paragraph(text)

        i += 1

    doc.save(output_file)
    sys.stdout.buffer.write('[OK] Created: {}\n'.format(output_file).encode('utf-8'))

if __name__ == '__main__':
    try:
        parse_markdown_to_docx('기술아키텍처.md', '기술아키텍처.docx')
    except Exception as e:
        sys.stdout.buffer.write('[ERROR] {}\n'.format(str(e)).encode('utf-8'))
        import traceback
        traceback.print_exc()
